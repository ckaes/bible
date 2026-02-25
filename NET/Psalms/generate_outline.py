"""
Generate Psalms_Outline.docx — a short outline listing psalms by category name only,
no verse text. Should be only a few pages.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

CATEGORIES = [
    {
        "name": "Hymns of Praise",
        "psalms": [8, 24, 29, 33, 47, 48, 50, 68, 76, 81, 87, 93, 95, 96, 97, 98, 99,
                   100, 103, 104, 105, 111, 113, 114, 115, 117, 135, 136, 145, 146,
                   147, 148, 149, 150],
    },
    {
        "name": "Laments and Cries for Help",
        "psalms": [3, 4, 5, 7, 10, 12, 13, 14, 17, 22, 25, 26, 28, 31, 35, 36, 39, 41,
                   42, 43, 44, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 64, 69, 70, 71,
                   74, 77, 79, 80, 82, 83, 85, 86, 88, 90, 94, 106, 108, 109, 137, 139,
                   140, 141, 142],
    },
    {
        "name": "Songs of Trust and Confidence",
        "psalms": [11, 16, 23, 27, 46, 62, 63, 84, 91],
    },
    {
        "name": "Penitential Psalms",
        "psalms": [6, 32, 38, 51, 102, 143],
    },
    {
        "name": "Thanksgiving and Testimony",
        "psalms": [9, 18, 30, 34, 40, 65, 66, 67, 75, 92, 107, 116, 118, 138, 144],
    },
    {
        "name": "Royal and Messianic Psalms",
        "psalms": [2, 20, 21, 45, 72, 89, 101, 110],
    },
    {
        "name": "Wisdom and Torah",
        "psalms": [1, 15, 19, 37, 49, 73, 78, 112, 119],
    },
    {
        "name": "Songs of Ascent",
        "psalms": [120, 121, 122, 123, 124, 125, 126, 127, 128, 129, 130, 131, 132, 133, 134],
    },
]

doc = Document()

# Title
title_para = doc.add_heading("The Psalms: A Thematic Arrangement — Outline", level=0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

for cat in CATEGORIES:
    doc.add_heading(cat["name"], level=1)
    for psalm_num in sorted(cat["psalms"]):
        doc.add_paragraph(f"Psalm {psalm_num}", style="List Bullet")

out_path = Path(__file__).parent / "Psalms_Outline.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
