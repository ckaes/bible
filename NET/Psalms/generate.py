"""
Generate Psalms_Categorized.docx — 150 Psalms grouped into 8 thematic categories.
Uses psalms.json (sibling file) as the source of NET verse text.
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------

CATEGORIES = [
    {
        "name": "Hymns of Praise",
        "psalms": [8, 24, 29, 33, 47, 48, 50, 68, 76, 81, 87, 93, 95, 96, 97, 98, 99,
                   100, 103, 104, 105, 111, 113, 114, 115, 117, 135, 136, 145, 146,
                   147, 148, 149, 150],
        "intro": (
            "These psalms burst forth in pure, celebratory adoration of God's majesty, "
            "power, and faithfulness, unburdened by petition or complaint. They extol the "
            "Creator of heaven and earth, the Lord who governs history and reigns over all "
            "nations, inviting Israel and all creation to join in worship. Whether focused "
            "on God's work in creation, his acts in history, or his cosmic kingship, they "
            "model the Psalter's ultimate purpose: unqualified praise."
        ),
    },
    {
        "name": "Laments and Cries for Help",
        "psalms": [3, 4, 5, 7, 10, 12, 13, 14, 17, 22, 25, 26, 28, 31, 35, 36, 39, 41,
                   42, 43, 44, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 64, 69, 70, 71,
                   74, 77, 79, 80, 82, 83, 85, 86, 88, 90, 94, 106, 108, 109, 137, 139,
                   140, 141, 142],
        "intro": (
            "The largest collection in the Psalter, these psalms give voice to suffering—"
            "from enemies, illness, national catastrophe, or the silence of God—laying the "
            "human condition bare before the divine. Whether the cry is personal or communal, "
            "the lament genre moves with characteristic honesty: complaint, appeal, a statement "
            "of trust, and often a vow to praise God once deliverance comes, preserving faith "
            "precisely by refusing to suppress pain."
        ),
    },
    {
        "name": "Songs of Trust and Confidence",
        "psalms": [11, 16, 23, 27, 46, 62, 63, 84, 91],
        "intro": (
            "Where laments describe the struggle toward faith, these psalms express its fruit: "
            "a settled serenity in God's care that transcends present circumstances. Composed in "
            "the voice of one who has already found shelter, they affirm God as shepherd, refuge, "
            "light, and fortress—not because trouble has vanished, but because God's presence is "
            "more real than the trouble."
        ),
    },
    {
        "name": "Penitential Psalms",
        "psalms": [6, 32, 38, 51, 102, 143],
        "intro": (
            "These psalms confront sin directly, moving through honest confession of guilt, deep "
            "contrition before a holy God, and urgent plea for mercy, toward the assurance that "
            "forgiveness restores the broken relationship. Recognized since ancient times as a "
            "distinct and precious group, they teach that the path back to God requires neither "
            "self-justification nor despair, but the simple acknowledgment: \u2018I have sinned; "
            "be merciful to me.\u2019"
        ),
    },
    {
        "name": "Thanksgiving and Testimony",
        "psalms": [9, 18, 30, 34, 40, 65, 66, 67, 75, 92, 107, 116, 118, 138, 144],
        "intro": (
            "Sung in the aftermath of God's specific intervention, these psalms recount the "
            "crisis, the cry for help, and the deliverance received, then invite the wider "
            "community to witness and join in gratitude. They serve as public testimony that "
            "God hears and answers prayer, transforming personal experience of salvation into "
            "corporate worship."
        ),
    },
    {
        "name": "Royal and Messianic Psalms",
        "psalms": [2, 20, 21, 45, 72, 89, 101, 110],
        "intro": (
            "These psalms orbit around the Davidic king—his coronation, his covenant with God, "
            "his battles, and his unique role as the Lord's anointed. Though rooted in Israel's "
            "royal theology, they consistently outrun any historical king in their promises, and "
            "the New Testament reads them as pointing ultimately to the Messiah, in whom all that "
            "was hoped for the ideal king is finally fulfilled."
        ),
    },
    {
        "name": "Wisdom and Torah",
        "psalms": [1, 15, 19, 37, 49, 73, 78, 112, 119],
        "intro": (
            "These psalms breathe the air of Israel's wisdom tradition, meditating on righteous "
            "living, the mystery of suffering, the contrast between the righteous and the wicked, "
            "and above all the life-giving perfection of God's Torah. Whether teaching, warning, "
            "or meditating at length on God's instruction, they insist that wisdom is not abstract "
            "philosophy but the practical daily orientation of the whole life toward God."
        ),
    },
    {
        "name": "Songs of Ascent",
        "psalms": [120, 121, 122, 123, 124, 125, 126, 127, 128, 129, 130, 131, 132, 133, 134],
        "intro": (
            "Comprising the collection gathered under the heading \u2018A Song of Ascents\u2019 "
            "(Psalms 120\u2013134), these brief, intimate pieces were sung by pilgrims journeying "
            "to Jerusalem for the great festivals. They trace a spiritual arc from distress in "
            "distant lands to joyful arrival at the house of God, celebrating the blessing of "
            "family, the peace of Zion, and the sure hope that flows from waiting on the Lord "
            "in his holy city."
        ),
    },
]

# ---------------------------------------------------------------------------
# Load verse data
# ---------------------------------------------------------------------------

here = Path(__file__).parent
with open(here / "psalms.json", encoding="utf-8") as f:
    raw = json.load(f)

# Build lookup: {psalm_number: {verse_number: text}}
psalm_verses: dict[int, dict[int, str]] = {}
for entry in raw:
    ch = int(entry["chapter"])
    vs = int(entry["verse"])
    psalm_verses.setdefault(ch, {})[vs] = entry["text"]

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# --- Title ---
title_para = doc.add_heading("The Psalms: A Thematic Arrangement (NET Translation)", level=0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

total_verses = 0

for cat in CATEGORIES:
    # Category heading (Heading 1)
    doc.add_heading(cat["name"], level=1)

    # Intro paragraph
    doc.add_paragraph(cat["intro"])

    for psalm_num in sorted(cat["psalms"]):
        # Psalm heading (Heading 2)
        doc.add_heading(f"Psalm {psalm_num}", level=2)

        verses = psalm_verses.get(psalm_num, {})
        for verse_num in sorted(verses.keys()):
            verse_text = verses[verse_num].strip()
            para = doc.add_paragraph()
            # Bold verse reference
            run_ref = para.add_run(f"{psalm_num}:{verse_num}  ")
            run_ref.bold = True
            # Verse text
            para.add_run(verse_text)
            total_verses += 1

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

out_path = here / "Psalms_Categorized.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Total verses written: {total_verses}")
